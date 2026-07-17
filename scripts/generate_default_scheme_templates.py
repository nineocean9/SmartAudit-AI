from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT_DIR = Path("ruoyi-admin/src/main/resources/static/audit-template/default")


TEMPLATES = [
    {
        "filename": "economic-responsibility-audit-plan-template.docx",
        "title": "经济责任审计实施方案模板（Word版）",
        "type": "经济责任审计",
        "sections": [
            ("一、项目基本情况", [
                "被审计领导干部任职单位基本情况：单位性质、组织架构、人员规模、资产资金规模、主要职责。",
                "任职期间及审计期间：列明任职起止时间、审计覆盖年度、必要的延伸追溯范围。",
                "前期审计及整改情况：梳理历史审计发现问题、整改落实情况和未完成事项。",
            ]),
            ("二、审计目标", [
                "评价领导干部任职期间经济责任履行情况。",
                "关注重大经济决策、财政财务收支、国有资产管理、内部控制建设和廉政风险。",
                "揭示重大风险隐患，促进规范权力运行和资产资金安全。",
            ]),
            ("三、审计范围与重点", [
                "重大经济决策制定、执行和效果情况。",
                "财政财务收支真实性、合法性和效益性。",
                "预算执行、政府采购、合同管理、资产配置处置及对外投资。",
                "以前年度审计、巡视巡察、监督检查问题整改情况。",
            ]),
            ("四、审计组织与分工", [
                "项目组长负责总体组织协调和重大事项把关。",
                "主审负责审计实施方案执行、问题定性、证据复核和报告组织。",
                "成员按资金、资产、项目、内控、数据分析等模块分工。",
            ]),
            ("五、实施步骤与时间安排", [
                "审前调查：收集资料、开展风险评估、完善审计重点。",
                "现场实施：执行审计程序、形成取证材料和审计工作底稿。",
                "报告阶段：汇总问题、征求意见、复核审理、出具报告。",
                "整改跟踪：建立问题清单、督促整改、评价整改成效。",
            ]),
            ("六、质量控制与风险提示", [
                "落实三级复核制度，重大问题及时请示报告。",
                "证据应充分、适当、关联，问题定性应准确。",
                "严格执行保密、廉政和现场纪律要求。",
            ]),
        ],
    },
    {
        "filename": "financial-revenue-expenditure-audit-plan-template.docx",
        "title": "财务收支审计实施方案模板（Word版）",
        "type": "财务收支审计",
        "sections": [
            ("一、项目基本情况", [
                "被审计单位财务管理体制、会计核算方式、账户设置和预算管理情况。",
                "审计期间收入、支出、资产、负债、净资产等总体规模。",
            ]),
            ("二、审计目标", [
                "核实财务收支真实性、合法性、完整性。",
                "评价预算执行、资金使用绩效和内部控制有效性。",
            ]),
            ("三、审计范围与重点", [
                "收入确认、票据管理、非税收入和往来款项。",
                "支出审批、报销合规、津补贴发放、三公经费和会议培训费。",
                "预算编制执行、结转结余、专项资金管理。",
                "资产采购、登记、使用、处置及账实相符情况。",
            ]),
            ("四、审计方法", [
                "账表、账账、账实核对。",
                "凭证抽查、银行流水核验、函证及访谈。",
                "利用数据分析识别异常支付、重复报销和超预算支出。",
            ]),
            ("五、实施安排", [
                "准备阶段完成资料清单、风险清单和抽样方案。",
                "现场阶段完成穿行测试、控制测试和实质性测试。",
                "报告阶段完成问题复核、法规依据匹配和整改建议。",
            ]),
            ("六、质量控制", [
                "重点关注资金安全、审批链条完整、附件真实有效。",
                "审计证据应覆盖金额、事实、责任主体和适用依据。",
            ]),
        ],
    },
    {
        "filename": "special-audit-plan-template.docx",
        "title": "专项审计实施方案模板（Word版）",
        "type": "专项审计",
        "sections": [
            ("一、专项背景", [
                "说明专项资金、专项政策或专项事项来源、管理要求和实施周期。",
                "列明主管部门、实施单位、受益对象及资金流转链条。",
            ]),
            ("二、审计目标", [
                "检查专项政策落实、资金分配使用、项目实施和绩效目标完成情况。",
                "揭示资金闲置、挤占挪用、虚报冒领、绩效不达标等问题。",
            ]),
            ("三、审计范围与重点", [
                "资金申报、审批、拨付、使用、结算全过程。",
                "项目立项、实施进度、验收管理和成果应用。",
                "制度建设、职责分工、监督检查和信息公开。",
                "绩效目标设置、过程监控和结果评价。",
            ]),
            ("四、审计程序", [
                "梳理政策依据和资金台账，建立项目清单。",
                "开展资金流向追踪和项目现场核查。",
                "抽查合同、发票、验收资料和受益对象真实性。",
                "进行绩效指标比对和异常数据筛查。",
            ]),
            ("五、时间安排", [
                "审前准备、现场核查、问题确认、报告编制、整改跟踪分阶段实施。",
            ]),
            ("六、风险控制", [
                "关注政策口径变化、跨部门数据一致性和现场核查覆盖率。",
                "重大疑点应延伸核查并形成完整证据链。",
            ]),
        ],
    },
    {
        "filename": "engineering-audit-plan-template.docx",
        "title": "工程审计实施方案模板（Word版）",
        "type": "工程审计",
        "sections": [
            ("一、工程项目概况", [
                "项目名称、建设地点、建设规模、建设内容、投资概算和资金来源。",
                "建设单位、代建单位、设计、监理、施工、造价咨询等参建主体。",
                "项目立项、招投标、合同签订、施工进度、竣工验收和结算情况。",
            ]),
            ("二、审计目标", [
                "评价工程建设程序合规性、投资控制有效性和资金使用真实性。",
                "揭示招投标、合同管理、工程变更、签证、结算和质量管理中的风险。",
            ]),
            ("三、审计范围与重点", [
                "项目决策立项、概预算审批和建设程序履行。",
                "招投标文件、评标过程、中标结果和合同条款。",
                "工程变更、现场签证、隐蔽工程、材料设备采购。",
                "工程量计算、定额套用、费用计取和结算真实性。",
                "工程款支付、资金来源、财务核算和资产移交。",
            ]),
            ("四、审计方法", [
                "资料审查、合同比对、现场踏勘、工程量复核。",
                "利用造价软件、清单计价规则和市场价格信息进行复核。",
                "对重大变更和异常签证开展延伸核查。",
            ]),
            ("五、实施步骤", [
                "准备阶段收集项目资料并制定抽样复核计划。",
                "现场阶段完成踏勘、测量、访谈和证据固定。",
                "结论阶段完成造价复核、问题定性和整改建议。",
            ]),
            ("六、质量与风险控制", [
                "工程量复核应保存计算过程和依据。",
                "涉及专业判断事项应组织复核或专家咨询。",
                "现场证据应具备时间、地点、人员和影像记录。",
            ]),
        ],
    },
]

REPORT_TEMPLATE = {
    "filename": "audit-report-draft-template.docx",
    "title": "审计报告模板（Word版）",
    "type": "审计报告",
    "sections": [
        ("一、基本情况", ["说明审计项目来源、被审计单位基本情况、审计期间、审计范围和审计组织实施情况。"]),
        ("二、审计评价意见", ["围绕财务收支、资产管理、项目建设、内部控制、政策执行等方面形成总体评价。"]),
        ("三、审计发现的主要问题", ["按照问题事实、金额影响、法规依据、责任主体的结构逐项列示。"]),
        ("四、审计处理意见和建议", ["提出处理处罚意见、管理改进建议和风险防控措施。"]),
        ("五、整改要求", ["明确整改责任、整改期限、反馈方式和后续跟踪安排。"]),
    ],
}


def set_run_font(run, font_name="仿宋", size=12, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold


def build_template(item):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "仿宋"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    normal.font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(title.add_run(item["title"]), "宋体", 18, True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(meta.add_run(f"适用类型：{item['type']}    模板版本：V1.0"), "仿宋", 11)

    for heading, bodies in item["sections"]:
        paragraph = doc.add_paragraph()
        set_run_font(paragraph.add_run(heading), "黑体", 14, True)
        for index, text in enumerate(bodies, 1):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(24)
            paragraph.paragraph_format.line_spacing = 1.5
            set_run_font(paragraph.add_run(f"{index}. {text}"), "仿宋", 12)

    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    for index, header in enumerate(["阶段", "主要任务", "责任人", "完成时限"]):
        set_run_font(table.rows[0].cells[index].paragraphs[0].add_run(header), "黑体", 10.5, True)
    rows = [
        ["审前准备", "资料收集、风险评估、方案细化", "", ""],
        ["现场实施", "审计取证、底稿编制、问题沟通", "", ""],
        ["报告整改", "报告编制、征求意见、整改跟踪", "", ""],
    ]
    for row_index, row in enumerate(rows, 1):
        for cell_index, value in enumerate(row):
            set_run_font(table.rows[row_index].cells[cell_index].paragraphs[0].add_run(value), "仿宋", 10.5)

    return doc


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for template in TEMPLATES:
        target = OUT_DIR / template["filename"]
        build_template(template).save(target)
        print(target)
    report_target = OUT_DIR / REPORT_TEMPLATE["filename"]
    build_template(REPORT_TEMPLATE).save(report_target)
    print(report_target)


if __name__ == "__main__":
    main()

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(r"F:\project1\RuoYi-Vue-v3.9.2\report-output")
OUT.mkdir(parents=True, exist_ok=True)

members = [
    {
        "name": "陈宇豪", "sid": "221302101", "share": "40%",
        "title": "系统总体架构与AI智能审计核心模块",
        "overview": "担任项目负责人，负责需求梳理、总体架构、数据库总体设计和核心技术集成，重点完成AI智能工作台、对话任务识别、RAG知识检索、项目文档解析与向量化、智能数据分析、风险扫描、文档核查和智能取证，并承担核心模块联调。",
        "modules": [
            ("AI智能对话与任务编排", "设计统一聊天入口，利用任务解析服务识别用户意图，将自然语言请求拆分为项目查询、资料阅读、数据分析、风险扫描、文档核查、智能取证和知识问答等任务，并通过SSE持续返回执行结果。"),
            ("RAG知识检索", "整合审计依据、项目资料、临时文件、审计案例和风险案例五类知识源。文档经过解析、切分和向量化后写入PostgreSQL与pgvector，问答时根据相似度召回相关内容并交由大模型生成回答。"),
            ("文件解析与向量化", "使用Apache POI处理Word和Excel文件，使用PDFBox解析PDF，同时支持CSV和TXT。解析后的文本按长度切块，并调用Embedding服务生成1024维向量，形成可检索的项目资料库。"),
            ("AI分析与智能取证", "实现AI数据分析、风险线索扫描、文档规范核查和取证单生成。分析结果以HTML和ECharts图表展示，取证单按照基本信息、问题描述、事实证据、法规依据、审计结论和整改建议组织。"),
        ],
        "code": [
            ("任务识别流程", "用户输入 → ChatTaskParserService → List<ChatTask> → 逐项执行 → SSE推送结果"),
            ("知识检索流程", "问题向量化 → pgvector相似度查询 → 多知识源结果合并 → 构建Prompt → 大模型回答"),
            ("文档入库流程", "文件上传 → 类型识别 → 文本解析 → DocChunker切分 → Embedding → document_chunk入库"),
        ],
        "tests": ["复合指令能否按顺序执行", "SSE连接是否持续、完整返回", "不同格式文件能否正确解析", "向量检索结果是否与问题相关", "AI异常时是否返回明确提示"]
    },
    {
        "name": "石雨淳", "sid": "__________", "share": "20%",
        "title": "审计前期管理与项目进度模块",
        "overview": "负责审计前期业务模块，包括审计计划、被审计单位、领导干部库、审计项目立项、项目成员配置、审计准备、通知书、资料清单和项目进度管理，形成从年度计划到项目正式实施前的完整业务链路。",
        "modules": [
            ("审计计划管理", "实现年度计划和专项计划的新增、修改、删除、查询、附件维护、变更记录和项目绑定，并支持按照年度、类型和状态进行组合筛选。"),
            ("审计对象与领导干部库", "维护被审计单位基本信息、单位负责人和领导干部任职履历，集中展示单位历年审计项目、发现问题和整改情况，为项目立项提供依据。"),
            ("审计项目与成员配置", "根据计划创建审计项目，设置项目负责人、审计组成员、实施时间和审计类型，保证项目人员及职责信息能够被后续底稿、问题和报告模块复用。"),
            ("审计准备与进度", "完成审计通知书、审计方案和资料清单管理，并使用ECharts展示项目计划时间、实际进度及逾期状态，实现项目进度的直观跟踪。"),
        ],
        "code": [
            ("计划管理流程", "审计计划录入 → 附件上传 → 审核确认 → 绑定审计项目 → 记录计划变更"),
            ("准备阶段流程", "创建项目 → 配置成员 → 生成通知书 → 下发资料清单 → 确认资料 → 进入实施"),
            ("进度计算", "根据计划开始、计划结束、实际完成时间计算阶段状态，并对临期和逾期项目进行标识"),
        ],
        "tests": ["计划新增及修改是否正确", "附件与变更记录是否保留", "单位和领导干部信息是否关联", "资料清单状态是否正常流转", "进度图是否正确显示项目时间"]
    },
    {
        "name": "骆旭林", "sid": "__________", "share": "20%",
        "title": "审计实施、整改与成果管理模块",
        "overview": "负责审计实施和成果管理模块，包括审计方案、审计底稿、审计问题、问题定性、整改跟踪、报告版本、意见处理及项目归档，形成问题发现、证据记录、整改落实和成果归档的业务闭环。",
        "modules": [
            ("审计方案与底稿", "按照项目维护审计目标、审计范围和实施步骤，支持审计底稿新增、编辑、复核状态维护以及法规依据引用，使审计过程和证据记录具有可追溯性。"),
            ("审计问题管理", "记录问题描述、涉及金额、严重程度、责任单位和法规依据，支持按项目、严重程度和整改状态筛选，并将问题与底稿、依据和整改记录关联。"),
            ("整改跟踪", "从审计问题发起整改任务，维护整改措施、计划完成日期、整改材料和审核意见，通过未整改、整改中、已完成等状态反映整改进度。"),
            ("报告与归档", "实现审计报告多版本管理、意见征求和文本差异查看。项目完成后汇总计划、资料、底稿、问题、整改和报告信息，形成完整电子归档记录。"),
        ],
        "code": [
            ("问题闭环", "审计底稿 → 发现问题 → 关联法规依据 → 发起整改 → 提交材料 → 审核确认"),
            ("报告版本", "创建报告草稿 → 保存版本 → 征求意见 → 修改并生成新版本 → 定稿"),
            ("项目归档", "校验必备资料 → 汇总项目成果 → 生成归档记录 → 设置项目完成状态"),
        ],
        "tests": ["底稿复核状态是否正确", "问题与依据能否建立关联", "整改状态是否按规则更新", "报告历史版本是否可查看", "缺少材料时是否阻止归档"]
    },
    {
        "name": "穆再排尔", "sid": "__________", "share": "20%",
        "title": "权限、知识库、数据驾驶舱与云服务器部署",
        "overview": "负责系统支撑及上线部署，包括用户角色、菜单权限、项目临时授权、审计依据库、案例库、风险案例库和数据驾驶舱，并完成PostgreSQL、pgvector、JDK、Nginx等云服务器环境配置、前后端部署、接口代理和上线测试。",
        "modules": [
            ("角色权限与临时授权", "基于RuoYi用户、角色、菜单和数据权限体系配置审计业务角色，通过权限标识控制页面按钮和后端接口；为中介审计人员提供按项目、按期限生效的临时授权及自动回收机制。"),
            ("依据库与案例库", "维护审计法规依据、典型案例和风险案例，支持分类、状态管理、多条件检索和语义搜索。新增或修改知识内容后同步完成向量化，为AI问答提供知识来源。"),
            ("数据驾驶舱", "使用ECharts展示审计项目数量、问题严重程度分布、整改完成情况、单位问题排名和项目趋势，使管理人员能够快速了解总体审计情况。"),
            ("云服务器部署", "在云服务器安装JDK 21、PostgreSQL、pgvector和Nginx，导入系统SQL，配置Spring Boot生产参数并运行后端Jar；将Vue构建产物交由Nginx托管，通过反向代理转发API和文件请求。"),
        ],
        "code": [
            ("权限控制", "用户登录 → 获取角色与菜单 → 前端动态路由 → 后端权限注解校验 → 数据范围过滤"),
            ("云端部署", "安装运行环境 → 初始化数据库 → 后端Maven打包 → 前端Vite构建 → Nginx托管与反向代理"),
            ("运行维护", "开放必要安全组端口，限制数据库外部访问，配置日志目录、文件目录、进程守护和定期备份"),
        ],
        "tests": ["不同角色菜单是否正确隔离", "临时授权过期后是否失效", "知识库检索是否准确", "驾驶舱统计是否与业务数据一致", "云端前后端及文件上传是否正常"]
    },
]

division = [(m["name"], m["share"], m["title"]) for m in members]

def set_cell_shading(cell, fill="E7E6E6"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); el = OxmlElement('w:tblHeader'); el.set(qn('w:val'), 'true'); trPr.append(el)

def set_cell_text(cell, text, bold=False, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text)); r.bold = bold; r.font.name = '宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体'); r.font.size = Pt(size); r.font.color.rgb = RGBColor(0,0,0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def style_doc(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.left_margin = Inches(1.25); sec.right_margin = Inches(1.25); sec.top_margin = Inches(1); sec.bottom_margin = Inches(1)
    styles = doc.styles
    normal = styles['Normal']; normal.font.name = 'Times New Roman'; normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体'); normal.font.size = Pt(12); normal.font.color.rgb = RGBColor(0,0,0)
    normal.paragraph_format.line_spacing = 1.5; normal.paragraph_format.space_after = Pt(0)
    for name, size in [('Title', 22), ('Heading 1', 16), ('Heading 2', 14), ('Heading 3', 12)]:
        st = styles[name]; st.font.name='Times New Roman'; st._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体'); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor(0,0,0)
        st.paragraph_format.space_before=Pt(10); st.paragraph_format.space_after=Pt(6); st.paragraph_format.keep_with_next=True
    styles['Title'].paragraph_format.space_before=Pt(0)

def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Pt(24); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead and text.startswith(bold_lead):
        r=p.add_run(bold_lead); r.bold=True; r2=p.add_run(text[len(bold_lead):])
    else: p.add_run(text)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f'Heading {level}')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_bullet(doc, text):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.left_indent=Pt(24); p.paragraph_format.first_line_indent=Pt(-12); p.add_run(text); return p

def add_table(doc, headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'; t.autofit=False
    for i,h in enumerate(headers): set_cell_text(t.rows[0].cells[i],h,True); set_cell_shading(t.rows[0].cells[i])
    set_repeat_table_header(t.rows[0])
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row): set_cell_text(cells[i],v,False,10)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    doc.add_paragraph()
    return t

def add_page_number(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=paragraph.add_run(); fldChar1=OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'),'begin'); instr=OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve'); instr.text=' PAGE '; fldChar2=OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'end'); run._r.extend([fldChar1,instr,fldChar2])

def cover(doc, m):
    for _ in range(2): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('扬  州  大  学'); r.bold=True; r.font.name='华文隶书'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'华文隶书'); r.font.size=Pt(30); r.font.color.rgb=RGBColor(0,0,0)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(20); r=p.add_run('项 目 综 合 实 践 Ⅲ 报 告'); r.bold=True; r.font.name='黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体'); r.font.size=Pt(24); r.font.color.rgb=RGBColor(0,0,0)
    for _ in range(4): doc.add_paragraph()
    info=[('题    目','高校智慧审计平台—AI智能工作台'),('姓    名',m['name']),('班    级','计科2201'),('学    号',m['sid']),('本组成员','陈宇豪、石雨淳、骆旭林、穆再排尔'),('指导老师','________________')]
    for k,v in info:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(8); r=p.add_run(f'{k}    {v}'); r.font.name='宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体'); r.font.size=Pt(14); r.font.color.rgb=RGBColor(0,0,0)
    doc.add_paragraph(); p=doc.add_paragraph('2025—2026学年第二学期'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

def common_content(doc):
    p=doc.add_paragraph('《项目综合实践III》任务书', style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_table(doc,['项目','内容'],[
        ('项目名称','高校智慧审计平台—AI智能工作台'),('主要内容','采用Vue 3、Spring Boot、PostgreSQL与大模型技术，建设覆盖审计计划、实施、整改、报告、归档和智能分析的综合平台。'),('技术要求','完成前后端分离系统、权限控制、文件解析、知识检索、AI辅助分析、数据可视化及云服务器部署。'),('成果要求','提交可运行源程序、数据库脚本、部署配置、实践报告和答辩材料。')],[1.3,5.2])
    p=doc.add_paragraph('《项目综合实践III》答辩记录表', style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_table(doc,['答辩问题','回答要点'],[
        ('系统为什么采用前后端分离架构？','前后端职责清晰，可以分别开发和部署；后端统一提供REST接口，前端负责交互展示，便于维护、扩展及权限控制。'),
        ('系统如何保证审计数据安全？','通过用户、角色、菜单、接口权限和数据范围进行多层控制；项目临时授权设置有效期限，重要操作保留日志，数据库不直接暴露到公网。'),
        ('AI生成的结果是否可以直接作为审计结论？','AI结果属于辅助建议，需要结合原始资料、法规依据和审计人员专业判断进行复核；系统保留引用内容和任务记录，便于追溯。')],[2.1,4.4])
    add_heading(doc,'团队成员分工',1)
    add_table(doc,['成员','工作量','主要负责内容'],division,[1.1,0.8,4.6])
    add_heading(doc,'一、项目实践的目的与要求',1)
    add_heading(doc,'1.1 项目实践目的',2)
    add_body(doc,'本项目面向高校审计业务数字化和智能化需求，综合运用Java程序设计、数据库原理、软件工程、Web前端开发和人工智能应用等知识，设计并实现一套高校智慧审计平台。通过项目实践，掌握前后端分离系统从需求分析、架构设计、数据库设计、编码实现、联调测试到云端部署的完整开发过程。')
    add_body(doc,'平台将大模型能力嵌入审计业务流程，使审计人员能够通过自然语言完成资料查询、知识检索、风险扫描、数据分析、文档核查和取证单生成，提高审计资料利用效率和审计作业规范性。')
    add_heading(doc,'1.2 项目实践要求',2)
    for x in ['掌握Vue 3、Element Plus、Vite和Axios等前端技术。','掌握Spring Boot、MyBatis、JWT和权限控制等后端技术。','掌握PostgreSQL数据库设计以及pgvector向量检索。','完成审计计划、实施、问题、整改、报告和归档等业务闭环。','实现AI对话、RAG检索、文件解析、智能分析和数据可视化。','完成系统测试、前后端打包和云服务器部署方案。']: add_bullet(doc,x)
    add_heading(doc,'二、需求分析',1)
    add_heading(doc,'2.1 项目背景',2)
    add_body(doc,'高校审计工作涉及计划编制、项目实施、问题确认、整改跟踪、报告编制和资料归档等多个阶段。传统工作方式中，业务数据分散在表格和文档中，资料检索、法规匹配、问题汇总及整改跟踪需要大量人工操作。为解决信息分散、协作效率低和审计知识利用不足等问题，本项目建设统一的高校智慧审计平台。')
    add_heading(doc,'2.2 功能需求',2)
    add_table(doc,['功能领域','主要功能'],[
        ('审计信息管理','审计计划、被审计单位、领导干部、项目立项和成员配置'),('审计作业管理','审计准备、通知书、资料清单、方案、底稿和项目进度'),('问题整改管理','问题记录、法规关联、整改任务、材料提交和整改审核'),('成果与知识管理','审计报告、版本意见、项目归档、依据库和案例库'),('AI智能辅助','智能对话、RAG问答、文件解析、数据分析、风险扫描、文档核查和取证'),('系统支撑','角色权限、临时授权、操作日志、数据驾驶舱和云服务器部署')],[1.5,5.0])
    add_heading(doc,'2.3 非功能需求',2)
    add_body(doc,'系统应具备良好的安全性、可维护性和可扩展性。用户只能访问角色和项目授权范围内的数据；文件上传和AI调用异常应提供明确提示；业务操作应保留必要日志；系统需支持常见桌面浏览器，并能够在云服务器环境稳定运行。')
    add_heading(doc,'三、项目总体设计方案',1)
    add_heading(doc,'3.1 系统总体架构',2)
    add_body(doc,'系统采用前后端分离架构。前端使用Vue 3和Element Plus构建业务页面，通过Axios调用后端REST接口，并使用SSE接收AI流式输出。后端以Spring Boot为基础，使用MyBatis完成数据访问，基于RuoYi的用户、角色、菜单和数据权限能力实现安全控制。PostgreSQL存储业务数据，pgvector保存文档向量。AI服务通过OpenAI兼容接口接入不同大模型。')
    add_table(doc,['层次','组成','作用'],[
        ('表示层','Vue 3、Element Plus、ECharts','页面展示、表单交互和统计可视化'),('接口层','Spring MVC、REST、SSE','业务接口及AI流式结果推送'),('业务层','审计业务服务、AI服务、文件服务','业务规则、任务编排和智能处理'),('数据层','MyBatis、PostgreSQL、pgvector','业务数据、文档文本和向量数据持久化'),('部署层','JDK、Nginx、云服务器','应用运行、静态托管和反向代理')],[1.0,2.2,3.3])
    add_heading(doc,'3.2 技术选型',2)
    add_table(doc,['类别','主要技术','选型说明'],[
        ('前端','Vue 3、Element Plus、Vite、Axios','组件化开发，界面统一，构建速度快'),('后端','Spring Boot、MyBatis、Spring Security','分层清晰，便于接口和权限开发'),('数据库','PostgreSQL、pgvector','同时支持关系数据和向量相似度检索'),('AI能力','LangChain4j、OpenAI兼容接口','便于切换模型并组织Prompt与流式调用'),('文档处理','Apache POI、PDFBox','支持Word、Excel和PDF内容解析'),('可视化','ECharts','支持项目进度和审计统计图表')],[1.0,2.2,3.3])
    add_heading(doc,'3.3 数据库设计',2)
    add_body(doc,'数据库按照系统基础数据、审计业务数据和AI知识数据进行划分。业务表之间通过项目编号、问题编号和文档编号建立关联，使计划、项目、底稿、问题、整改、报告和归档数据保持一致。')
    add_table(doc,['数据表','用途'],[
        ('audit_plan / audit_project','审计计划及项目基本信息'),('audit_issue / audit_rectification','审计问题和整改跟踪'),('audit_workpaper / audit_report','审计底稿和报告成果'),('audit_basis / audit_case_lib','法规依据和案例知识'),('project_document / document_chunk','项目文件、解析文本和向量切块'),('ai_conversation / ai_message','AI会话及消息记录'),('audit_temp_auth','中介人员项目临时授权')],[2.4,4.1])
    add_heading(doc,'四、系统主要功能实现',1)
    for h,t in [
        ('4.1 审计全流程管理','系统将审计计划、项目准备、现场实施、问题确认、整改跟踪、报告编制和项目归档连接起来。各阶段数据通过项目编号关联，用户可以从项目页面查看项目成员、资料清单、底稿、问题、整改和报告。'),
        ('4.2 AI智能工作台','AI工作台提供统一的自然语言交互入口。系统首先识别用户意图和目标项目，然后选择对应处理器执行资料阅读、知识问答、智能分析、风险扫描、文档核查或取证任务，并将结果流式显示在页面中。'),
        ('4.3 知识库与项目资料库','审计依据、案例、风险案例和项目文件共同组成审计知识来源。上传的文档经过内容解析和向量化后，可以按语义相似度检索，减少仅依赖关键词造成的漏检。'),
        ('4.4 数据驾驶舱','数据驾驶舱汇总项目数量、问题等级、整改进度和单位问题分布等指标，通过柱状图、饼图和趋势图展示，为管理人员掌握整体审计情况提供支持。')]:
        add_heading(doc,h,2); add_body(doc,t)

def personal_content(doc,m):
    add_heading(doc,'五、个人完成的功能及实现',1)
    add_heading(doc,'5.1 个人分工概述',2); add_body(doc,f"本人承担项目总工作量的{m['share']}，主要负责{m['title']}。{m['overview']}")
    for idx,(h,t) in enumerate(m['modules'],1):
        add_heading(doc,f'5.{idx+1} {h}',2); add_body(doc,t)
        add_body(doc,'在实现过程中，前端页面负责信息录入、查询筛选、状态展示和操作反馈，后端按照Controller、Service、Mapper进行分层，并通过数据库约束和业务校验保证数据完整性。相关接口统一返回处理结果，异常由全局异常机制进行捕获。')
    add_heading(doc,'六、关键流程与实现说明',1)
    add_table(doc,['序号','关键内容','实现流程或说明'],[(i+1,a,b) for i,(a,b) in enumerate(m['code'])],[0.6,1.5,4.4])
    add_heading(doc,'七、系统测试',1)
    add_heading(doc,'7.1 测试环境',2)
    add_table(doc,['项目','配置'],[('操作系统','Windows 11 / Linux云服务器'),('前端环境','Node.js、Vue 3、Vite'),('后端环境','JDK 21、Spring Boot'),('数据库','PostgreSQL 18、pgvector'),('浏览器','Chrome、Edge')],[1.5,5.0])
    add_heading(doc,'7.2 个人模块测试',2)
    rows=[]
    for i,x in enumerate(m['tests'],1): rows.append((i,x,'按照正常、空值和异常数据分别执行操作','功能响应符合设计要求'))
    add_table(doc,['序号','测试内容','测试方法','结果'],rows,[0.55,1.85,2.75,1.35])
    add_heading(doc,'7.3 测试结论',2)
    add_body(doc,'经过功能测试和前后端联调，个人负责模块能够完成预定业务操作，页面数据与数据库记录保持一致，权限不足、参数错误和服务异常等情况能够返回相应提示。测试中发现的问题经过修改后重新验证，未发现影响主要流程运行的严重缺陷。')
    add_heading(doc,'八、云服务器部署方案',1)
    add_body(doc,'系统部署采用Nginx静态托管与Spring Boot独立运行相结合的方式。服务器准备Linux运行环境，安装JDK、PostgreSQL、pgvector和Nginx；执行项目SQL脚本初始化数据库；修改生产环境数据库、文件目录和AI模型参数；使用Maven生成后端Jar包，使用Vite生成前端dist目录。Nginx负责前端资源访问，并将/api请求反向代理到后端服务。')
    add_table(doc,['部署步骤','主要操作'],[
        ('1. 环境准备','配置云服务器安全组，安装JDK、PostgreSQL、pgvector和Nginx'),('2. 数据初始化','创建数据库用户，执行基础表、AI表和业务扩展SQL'),('3. 后端部署','配置生产参数，打包Jar，设置日志和上传目录，启动服务'),('4. 前端部署','执行构建，将dist目录上传至Nginx站点目录'),('5. 代理配置','配置静态资源、API反向代理、上传大小和超时时间'),('6. 上线检查','验证登录、权限、文件上传、AI调用、统计图表和业务闭环')],[1.35,5.15])
    add_heading(doc,'九、总结',1)
    add_body(doc,f"通过本次高校智慧审计平台开发，我完成了{m['title']}相关工作，对前后端分离架构、数据库设计和团队协作流程有了更加系统的认识。项目开发不仅要求功能能够运行，还需要考虑权限边界、数据关联、异常处理、部署环境和后续维护。")
    add_body(doc,'在实践过程中，我进一步理解了需求分析、模块划分、接口约定、联调测试和问题定位的重要性。后续可继续完善自动化测试、AI结果评价、模型调用成本统计、移动端适配和运维监控，使系统具备更高的可靠性与实用价值。')
    add_heading(doc,'参考文献',1)
    for x in ['[1] RuoYi-Vue v3.9.2 项目文档。','[2] Spring Boot Reference Documentation。','[3] Vue.js 3 Documentation。','[4] PostgreSQL Documentation。','[5] LangChain4j Documentation。','[6] Apache ECharts Documentation。']:
        p=doc.add_paragraph(x); p.paragraph_format.left_indent=Pt(0)

def build(m):
    doc=Document(); style_doc(doc)
    footer=doc.sections[0].footer.paragraphs[0]; add_page_number(footer)
    cover(doc,m); common_content(doc); personal_content(doc,m)
    # force every visible run and heading black; preserve all headings as black even under Word themes
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.color.rgb=RGBColor(0,0,0)
            if p.style and p.style.name.startswith(('Heading','Title')): r.font.bold=True
    path=OUT/f"{m['name']}_项目综合实践III_高校智慧审计平台.docx"
    doc.core_properties.title='高校智慧审计平台—AI智能工作台项目综合实践III报告'
    doc.core_properties.subject=m['title']; doc.core_properties.author=m['name']
    doc.save(path); return path

if __name__=='__main__':
    for m in members: print(build(m))

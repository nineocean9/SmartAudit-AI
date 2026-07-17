from pathlib import Path
from zipfile import ZipFile
from docx import Document
from docx.oxml.ns import qn

root=Path(__file__).parent
expected={
    '陈宇豪':['AI智能对话与任务编排','RAG知识检索','文件解析与向量化','AI分析与智能取证'],
    '石雨淳':['审计计划管理','审计对象与领导干部库','审计项目与成员配置','审计准备与进度'],
    '骆旭林':['审计方案与底稿','审计问题管理','整改跟踪','报告与归档'],
    '穆再排尔':['角色权限与临时授权','依据库与案例库','数据驾驶舱','云服务器部署'],
}
common=['《项目综合实践III》任务书','《项目综合实践III》答辩记录表','团队成员分工','一、项目实践的目的与要求','二、需求分析','三、项目总体设计方案','四、系统主要功能实现','五、个人完成的功能及实现','六、关键流程与实现说明','七、系统测试','八、云服务器部署方案','九、总结','参考文献']
failed=False
for name,own in expected.items():
    path=next(root.glob(f'{name}_*.docx'))
    with ZipFile(path) as z:
        assert z.testzip() is None
    d=Document(path)
    text='\n'.join(p.text for p in d.paragraphs)+'\n'+'\n'.join(c.text for t in d.tables for row in t.rows for c in row.cells)
    missing=[x for x in common+own if x not in text]
    bad=[]
    for p in d.paragraphs:
        if p.style and (p.style.name.startswith('Heading') or p.style.name=='Title'):
            for r in p.runs:
                color=r._r.get_or_add_rPr().find(qn('w:color'))
                val=color.get(qn('w:val')) if color is not None else None
                if val not in (None,'000000','auto','AUTO'):
                    bad.append((p.text,val))
    ok=not missing and not bad and len(d.tables)==11 and len(d.sections)==1
    failed |= not ok
    print(f'{path.name}: ok={ok}, paragraphs={len(d.paragraphs)}, tables={len(d.tables)}, headings_black={not bad}, missing={missing}')
raise SystemExit(1 if failed else 0)
